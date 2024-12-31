
from flask import Flask, render_template, request, send_file
from docx import Document
import os

app = Flask(__name__)

UPLOAD_FOLDER = 'uploads'
OUTPUT_FOLDER = 'output'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/process_text', methods=['POST'])
def process_text():
    input_text = request.form['input_text']
    filename = generate_word_document(input_text.splitlines())
    return send_file(filename, as_attachment=True)

@app.route('/upload', methods=['POST'])
def upload_file():
    uploaded_file = request.files['file']
    if uploaded_file.filename != '':
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], uploaded_file.filename)
        uploaded_file.save(file_path)
        with open(file_path, 'r', encoding='utf-8') as f:
            input_text = f.readlines()
        filename = generate_word_document(input_text)
        return send_file(filename, as_attachment=True)

def generate_word_document(text_lines):
    doc = Document()
    doc.add_heading('בונה שערי נספחים - אליהו ביטון', level=1)
    for i, line in enumerate(text_lines, start=1):
        doc.add_page_break()
        doc.add_paragraph(f"נספח מספר {i}", style='Title')
        doc.add_paragraph(f"שם הנספח: {line.strip()}", style='Body Text')
    output_path = os.path.join(OUTPUT_FOLDER, 'נספחים.docx')
    doc.save(output_path)
    return output_path

if __name__ == '__main__':
    app.run(debug=True)
